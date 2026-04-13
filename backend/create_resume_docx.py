from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create document
doc = Document()

# Set default font
doc.styles['Normal'].font.name = 'Microsoft YaHei'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Title
title = doc.add_heading('率乔', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)

# Contact info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('22岁 | 18322738513 | 广州市番禺区 | 2507703520@qq.com')
contact_run.font.size = Pt(10)
contact_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Job intention
heading = doc.add_heading('求职意向', 1)
heading.runs[0].font.size = Pt(14)
heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

p = doc.add_paragraph()
p.add_run('AI应用工程师 / Agent工程师').font.size = Pt(11)

# Education
heading = doc.add_heading('教育背景', 1)
heading.runs[0].font.size = Pt(14)
heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

p = doc.add_paragraph()
p.add_run('武汉东湖学院').bold = True
p.add_run(' | ')
p.add_run('计算机科学与技术').bold = True
p.add_run(' | ')
p.add_run('本科').bold = True
tab_run = p.add_run()
tab_run._element.addprevious(doc._element.makeelement(qn('w:tab')))
p.add_run('2020.09 - 2024.06').font.color.rgb = RGBColor(100, 100, 100)

p = doc.add_paragraph('主修课程：Python编程基础、Java程序设计、数据结构、数据库原理、操作系统、计算机网络', style='List Bullet')
p.runs[0].font.size = Pt(10)

# Skills
heading = doc.add_heading('专业技能', 1)
heading.runs[0].font.size = Pt(14)
heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

# AI Skills
p = doc.add_paragraph()
p.add_run('AI/LLM相关').bold = True
p.runs[0].font.size = Pt(11)

skills_ai = [
    '熟练掌握大语言模型(LLM)应用开发，包括Prompt Engineering、RAG检索增强生成',
    '熟悉Multi-Agent架构设计，具备ReAct模式、任务编排、工具调用等实战经验',
    '掌握DeepSeek API、OpenAI API等主流大模型接口调用',
    '了解向量数据库(ChromaDB)和文本嵌入技术'
]
for skill in skills_ai:
    p = doc.add_paragraph(skill, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Backend Skills
p = doc.add_paragraph()
p.add_run('后端开发').bold = True
p.runs[0].font.size = Pt(11)

skills_backend = [
    '熟练掌握Python、Java、C#多语言开发',
    '熟练使用FastAPI、SpringBoot框架构建RESTful API',
    '熟悉MySQL关系型数据库设计与优化',
    '了解消息队列Kafka的使用及原理',
    '熟悉Linux环境部署与运维'
]
for skill in skills_backend:
    p = doc.add_paragraph(skill, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Frontend & Tools
p = doc.add_paragraph()
p.add_run('前端与工具').bold = True
p.runs[0].font.size = Pt(11)

skills_frontend = [
    '掌握React、TailwindCSS前端开发',
    '熟练使用Git版本控制、Docker容器化部署',
    '掌握IDEA、PyCharm、VS Code等开发工具'
]
for skill in skills_frontend:
    p = doc.add_paragraph(skill, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Certificates
p = doc.add_paragraph()
p.add_run('证书').bold = True
p.runs[0].font.size = Pt(11)

certs = [
    '工业和信息化人才专业知识测评证书（Java软件开发工程师科目）',
    'OceanBase数据库OBCA认证'
]
for cert in certs:
    p = doc.add_paragraph(cert, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Projects
heading = doc.add_heading('项目经历', 1)
heading.runs[0].font.size = Pt(14)
heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

# Project 1
p = doc.add_paragraph()
p.add_run('跨境智卖 - AI跨境电商运营助手').bold = True
p.runs[0].font.size = Pt(12)
p.add_run('\n')
p.add_run('独立开发 | 全栈开发').font.size = Pt(10)
tab_run = p.add_run()
tab_run._element.addprevious(doc._element.makeelement(qn('w:tab')))
p.add_run('2025.03 - 2025.04').font.color.rgb = RGBColor(100, 100, 100)
p.add_run('\n')

p = doc.add_paragraph()
p.add_run('项目简介：').bold = True
p.add_run('基于Multi-Agent架构的AI跨境电商智能运营平台，帮助卖家自动生成多语言商品文案、智能定价分析、竞品监控')

p = doc.add_paragraph()
p.add_run('技术栈：').bold = True
p.add_run('Python + FastAPI + React + DeepSeek API + ChromaDB + TailwindCSS')

p = doc.add_paragraph()
p.add_run('核心功能：').bold = True

features = [
    'AI商品文案工厂：支持Amazon/TikTok/Shopify多平台风格文案生成，支持6种语言一键生成',
    'RAG智能客服系统：基于ChromaDB构建知识库，实现向量检索+大模型生成的问答pipeline',
    'Multi-Agent定价分析系统：设计5个协作Agent（Planner/Retriever/Calculator/Generator/Checker），实现ReAct模式任务编排'
]
for feature in features:
    p = doc.add_paragraph(feature, style='List Bullet')
    p.runs[0].font.size = Pt(10)

p = doc.add_paragraph()
p.add_run('项目亮点：').bold = True

highlights = [
    '采用Multi-Agent架构，展示AI任务分解与协作能力',
    '实现完整的RAG pipeline，体现LLM应用开发能力',
    '前后端分离，支持Vercel+Render免费部署',
    'GitHub开源，包含完整技术文档'
]
for highlight in highlights:
    p = doc.add_paragraph(highlight, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Project 2
p = doc.add_paragraph()
p.add_run('苍穹外卖').bold = True
p.runs[0].font.size = Pt(12)
p.add_run('\n')
p.add_run('后端开发').font.size = Pt(10)
tab_run = p.add_run()
tab_run._element.addprevious(doc._element.makeelement(qn('w:tab')))
p.add_run('2022.12 - 2023.01').font.color.rgb = RGBColor(100, 100, 100)
p.add_run('\n')

p = doc.add_paragraph()
p.add_run('技术栈：').bold = True
p.add_run('SpringBoot + MyBatis-Plus + MySQL + JDK8')

p = doc.add_paragraph()
p.add_run('责任描述：').bold = True

descs = [
    '开发客户模块，实现商品展示、下单、购物车功能',
    '实现商品分类、搜索、推荐等展示功能',
    '完成订单生成、购物车增删改查等核心业务逻辑'
]
for desc in descs:
    p = doc.add_paragraph(desc, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Campus Experience
heading = doc.add_heading('在校经历', 1)
heading.runs[0].font.size = Pt(14)
heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

experiences = [
    '乒乓球校队成员（2020年加入）',
    '金融投资爱好者：4年股票实盘经验，阅读《价值》《金融心理学》《波浪理论》等书籍',
    '雪球平台作者：笔名"乔木嘎嘎"，发表个人复盘文章，总浏览量18.46万'
]
for exp in experiences:
    p = doc.add_paragraph(exp, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Self Evaluation
heading = doc.add_heading('自我评价', 1)
heading.runs[0].font.size = Pt(14)
heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

evals = [
    '积极认真，细心负责，具备较强的分析和问题解决能力',
    '对AI技术和LLM应用有浓厚兴趣，持续学习前沿技术',
    '具备全栈开发能力，能独立完成项目从设计到部署',
    '热爱编程，善于将业务需求转化为技术方案'
]
for ev in evals:
    p = doc.add_paragraph(ev, style='List Bullet')
    p.runs[0].font.size = Pt(10)

# Save document
output_path = r'C:\Users\qiao\Downloads\率乔_简历_AI应用工程师.docx'
doc.save(output_path)
print(f'Resume saved to: {output_path}')
