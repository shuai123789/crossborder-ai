from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create document
doc = Document()

# Set narrow margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# Set default font
doc.styles['Normal'].font.name = 'Microsoft YaHei'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
doc.styles['Normal'].font.size = Pt(10)

# Title
title = doc.add_heading('率乔', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)

# Contact info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('22岁 | 18322738513 | 广州市番禺区 | 2507703520@qq.com')
contact_run.font.size = Pt(9)
contact_run.font.color.rgb = RGBColor(80, 80, 80)

# Job intention
p = doc.add_paragraph()
p.add_run('求职意向：').bold = True
p.add_run('AI应用工程师 / Agent工程师')
p.paragraph_format.space_after = Pt(6)

# Education
p = doc.add_paragraph()
p.add_run('教育背景').bold = True
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

p = doc.add_paragraph()
p.add_run('武汉东湖学院').bold = True
p.add_run(' | 计算机科学与技术 | 本科 ')
p.add_run('2020.09 - 2024.06').font.color.rgb = RGBColor(100, 100, 100)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph('主修课程：Python编程基础、Java程序设计、数据结构、数据库原理、操作系统、计算机网络')
p.paragraph_format.space_after = Pt(6)

# Skills
p = doc.add_paragraph()
p.add_run('专业技能').bold = True
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

skills_text = '''AI/LLM：Prompt Engineering、RAG检索增强生成、Multi-Agent架构（ReAct模式）、DeepSeek API、向量数据库ChromaDB
后端：Python、Java、FastAPI、SpringBoot、MySQL、Kafka、Linux
前端：React、TailwindCSS、Git、Docker'''

p = doc.add_paragraph(skills_text)
p.paragraph_format.space_after = Pt(6)

# Certificates
p = doc.add_paragraph()
p.add_run('证书：').bold = True
p.add_run('工业和信息化人才专业知识测评证书（Java软件开发工程师）| OceanBase数据库OBCA认证')
p.paragraph_format.space_after = Pt(6)

# Projects
p = doc.add_paragraph()
p.add_run('项目经历').bold = True
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

# Project 1
p = doc.add_paragraph()
p.add_run('跨境智卖 - AI跨境电商运营助手').bold = True
p.add_run(' | 独立开发 ')
p.add_run('2025.03 - 2025.04').font.color.rgb = RGBColor(100, 100, 100)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('技术栈：Python + FastAPI + React + DeepSeek API + ChromaDB')
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('• AI商品文案工厂：支持Amazon/TikTok/Shopify多平台6种语言文案生成\n'
                     '• RAG智能客服：基于ChromaDB构建知识库，实现向量检索+大模型生成问答\n'
                     '• Multi-Agent定价分析：5个Agent协作（Planner/Retriever/Calculator/Generator/Checker），ReAct模式任务编排，自动生成定价策略报告')
p.paragraph_format.space_after = Pt(6)

# Project 2
p = doc.add_paragraph()
p.add_run('苍穹外卖').bold = True
p.add_run(' | 后端开发 ')
p.add_run('2022.12 - 2023.01').font.color.rgb = RGBColor(100, 100, 100)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('技术栈：SpringBoot + MyBatis-Plus + MySQL + JDK8')
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph('• 开发客户模块：商品展示（分类/搜索/推荐）、下单、购物车功能\n'
                     '• 实现订单生成、购物车增删改查等核心业务逻辑')
p.paragraph_format.space_after = Pt(6)

# Campus Experience
p = doc.add_paragraph()
p.add_run('在校经历').bold = True
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

p = doc.add_paragraph('• 乒乓球校队成员（2020年加入）\n'
                     '• 金融投资爱好者：4年股票实盘经验，阅读《价值》《金融心理学》《波浪理论》\n'
                     '• 雪球平台作者：笔名"乔木嘎嘎"，个人复盘文章总浏览量18.46万')
p.paragraph_format.space_after = Pt(6)

# Self Evaluation
p = doc.add_paragraph()
p.add_run('自我评价').bold = True
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

p = doc.add_paragraph('积极认真，细心负责，具备较强的分析和问题解决能力；对AI技术和LLM应用有浓厚兴趣，持续学习前沿技术；'
                     '具备全栈开发能力，能独立完成项目从设计到部署；热爱编程，善于将业务需求转化为技术方案。')

# Save document
output_path = r'C:\Users\qiao\Downloads\率乔_简历_AI应用工程师_紧凑版.docx'
doc.save(output_path)
print(f'Resume saved to: {output_path}')
